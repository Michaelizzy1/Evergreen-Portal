# results/result_docx.py
#
# Generates a formatted Word document (.docx) for a student's result.
# Called by the print_result view — returns an HttpResponse with the file.
#
# Install dependency:
#   pip install python-docx

from io import BytesIO
from docx import Document
from docx.shared import Pt, Inches, RGBColor, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_ALIGN_VERTICAL
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
from django.http import HttpResponse


# ── Colour constants ──────────────────────────────────────────────────────────
INK        = RGBColor(0x0d, 0x0f, 0x14)
GOLD       = RGBColor(0xC8, 0xA8, 0x4B)
GREEN      = RGBColor(0x1E, 0x7A, 0x44)
BLUE       = RGBColor(0x1E, 0x4F, 0xA3)
AMBER      = RGBColor(0xB8, 0x86, 0x0B)
RED        = RGBColor(0xC0, 0x39, 0x2B)
LIGHT_GREY = RGBColor(0xF0, 0xEE, 0xE8)
MID_GREY   = RGBColor(0xCC, 0xC8, 0xBE)
DARK_GREY  = RGBColor(0x3A, 0x3D, 0x47)
WHITE      = RGBColor(0xFF, 0xFF, 0xFF)


# ── XML helpers ───────────────────────────────────────────────────────────────

def hex_color(rgb: RGBColor) -> str:
    return f"{rgb[0]:02X}{rgb[1]:02X}{rgb[2]:02X}"


def set_cell_bg(cell, rgb: RGBColor):
    """Fill a table cell background."""
    tc   = cell._tc
    tcPr = tc.get_or_add_tcPr()
    shd  = OxmlElement('w:shd')
    shd.set(qn('w:val'),   'clear')
    shd.set(qn('w:color'), 'auto')
    shd.set(qn('w:fill'),  hex_color(rgb))
    tcPr.append(shd)


def set_cell_border(cell, **kwargs):
    """
    Set borders on a single cell.
    kwargs: top, bottom, left, right — each a dict with keys: sz, color, val
    """
    tc   = cell._tc
    tcPr = tc.get_or_add_tcPr()
    tcBorders = OxmlElement('w:tcBorders')
    for side in ('top', 'left', 'bottom', 'right'):
        cfg = kwargs.get(side)
        if cfg:
            el = OxmlElement(f'w:{side}')
            el.set(qn('w:val'),   cfg.get('val',   'single'))
            el.set(qn('w:sz'),    str(cfg.get('sz', 4)))
            el.set(qn('w:space'), '0')
            el.set(qn('w:color'), cfg.get('color', 'auto'))
            tcBorders.append(el)
    tcPr.append(tcBorders)


def cell_paragraph(cell, text, bold=False, size=10, color=INK,
                   align=WD_ALIGN_PARAGRAPH.LEFT, italic=False):
    """Clear a cell and add a single formatted paragraph."""
    cell.text = ''
    p   = cell.paragraphs[0]
    p.alignment = align
    run = p.add_run(str(text))
    run.bold   = bold
    run.italic = italic
    run.font.size  = Pt(size)
    run.font.color.rgb = color
    return p


def add_paragraph(doc, text='', bold=False, size=11, color=INK,
                  align=WD_ALIGN_PARAGRAPH.LEFT, space_before=0, space_after=6):
    p   = doc.add_paragraph()
    p.alignment = align
    p.paragraph_format.space_before = Pt(space_before)
    p.paragraph_format.space_after  = Pt(space_after)
    if text:
        run = p.add_run(text)
        run.bold  = bold
        run.font.size  = Pt(size)
        run.font.color.rgb = color
    return p


def add_horizontal_rule(doc, color=GOLD):
    """A thin coloured paragraph border as a divider."""
    p   = doc.add_paragraph()
    pPr = p._p.get_or_add_pPr()
    pBdr = OxmlElement('w:pBdr')
    bottom = OxmlElement('w:bottom')
    bottom.set(qn('w:val'),   'single')
    bottom.set(qn('w:sz'),    '6')
    bottom.set(qn('w:space'), '1')
    bottom.set(qn('w:color'), hex_color(color))
    pBdr.append(bottom)
    pPr.append(pBdr)
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.space_after  = Pt(6)


def grade_color(grade: str) -> RGBColor:
    return {'A': GREEN, 'B': BLUE, 'C': AMBER, 'D': RED, 'F': RED}.get(grade, INK)


# ── Main document builder ─────────────────────────────────────────────────────

def build_result_docx(student, term_cards) -> bytes:
    doc = Document()

    # ── Page setup: A4, 1.5cm margins ────────────────────────────────────────
    for section in doc.sections:
        section.page_width  = Cm(21)
        section.page_height = Cm(29.7)
        section.top_margin    = Cm(1.5)
        section.bottom_margin = Cm(1.5)
        section.left_margin   = Cm(1.8)
        section.right_margin  = Cm(1.8)

    # ── School header ─────────────────────────────────────────────────────────
    h = add_paragraph(doc, 'EVERGREEN ACADEMY', bold=True, size=18,
                      color=INK, align=WD_ALIGN_PARAGRAPH.CENTER,
                      space_before=0, space_after=2)
    add_paragraph(doc, 'Student Academic Result', bold=False, size=11,
                  color=DARK_GREY, align=WD_ALIGN_PARAGRAPH.CENTER,
                  space_before=0, space_after=2)
    add_horizontal_rule(doc)

    # ── Student bio block ─────────────────────────────────────────────────────
    # 2-column mini-table (no borders) for name / class / ID / session
    bio = doc.add_table(rows=2, cols=4)
    bio.style = 'Table Grid'

    labels = ['Student Name', 'Class', 'Student ID', 'Session']
    tc0    = term_cards[0] if term_cards else None
    values = [
        student.full_name,
        student.class_name,
        student.student_id,
        tc0['session'] if tc0 else '—',
    ]

    for i, (lbl, val) in enumerate(zip(labels, values)):
        lc = bio.rows[0].cells[i]
        vc = bio.rows[1].cells[i]
        cell_paragraph(lc, lbl, bold=True, size=8, color=DARK_GREY)
        cell_paragraph(vc, val, bold=True, size=11, color=INK)
        set_cell_bg(lc, LIGHT_GREY)
        # Remove all borders for a clean look
        no_border = {'val': 'none', 'sz': 0, 'color': 'FFFFFF'}
        for cell in (lc, vc):
            set_cell_border(cell,
                top=no_border, bottom=no_border,
                left=no_border, right=no_border)

    doc.add_paragraph()  # spacer

    # ── One section per term ──────────────────────────────────────────────────
    for idx, tc in enumerate(term_cards):

        if idx > 0:
            doc.add_page_break()

        # Term heading
        add_paragraph(doc, tc['term_label'].upper(), bold=True, size=13,
                      color=INK, space_before=4, space_after=2)
        add_horizontal_rule(doc, color=GOLD)

        # ── Summary stats row (4-cell table) ─────────────────────────────────
        stats = doc.add_table(rows=2, cols=4)
        stats.style = 'Table Grid'

        stat_labels = ['Total Score', 'Average', 'Class Position', 'Attendance']
        stat_values = [
            f"{tc['total_score']} / {tc['max_score']}",
            f"{tc['average']}%  (Grade {tc['overall_grade']})",
            f"{tc['position']} of {tc['total_students']}" if tc['position'] else '—',
            f"{tc['attendance']}%  ({tc['days_present']} present, {tc['days_absent']} absent)"
                if tc['attendance'] is not None else '—',
        ]

        col_w = Cm(4.4)
        for i, (lbl, val) in enumerate(zip(stat_labels, stat_values)):
            lc = stats.rows[0].cells[i]
            vc = stats.rows[1].cells[i]
            lc.width = col_w
            vc.width = col_w
            cell_paragraph(lc, lbl, bold=True, size=8, color=WHITE)
            cell_paragraph(vc, val, bold=True, size=10, color=INK)
            set_cell_bg(lc, INK)
            set_cell_bg(vc, LIGHT_GREY)
            thin = {'val': 'single', 'sz': 2, 'color': 'FFFFFF'}
            for cell in (lc, vc):
                set_cell_border(cell,
                    top=thin, bottom=thin, left=thin, right=thin)

        doc.add_paragraph()  # spacer

        # ── Highest subject ───────────────────────────────────────────────────
        if tc.get('highest_subject'):
            hs = tc['highest_subject']
            p  = doc.add_paragraph()
            p.paragraph_format.space_after = Pt(8)
            r1 = p.add_run('Highest Subject: ')
            r1.bold = True
            r1.font.size = Pt(10)
            r1.font.color.rgb = DARK_GREY
            r2 = p.add_run(f"{hs['name']}  —  {hs['total']} / 100")
            r2.bold = True
            r2.font.size = Pt(11)
            r2.font.color.rgb = GREEN

        # ── Subject results table ─────────────────────────────────────────────
        add_paragraph(doc, 'Subject Performance', bold=True, size=11,
                      color=INK, space_before=6, space_after=4)

        headers = ['Subject', 'CA (/40)', 'Exam (/60)', 'Total (/100)', 'Grade', 'Remark']
        # Column widths in DXA — must sum to content width
        # A4 with 1.8cm margins: usable ≈ 17.4cm
        col_widths_cm = [5.5, 2.2, 2.2, 2.8, 1.8, 2.9]

        tbl = doc.add_table(rows=1, cols=6)
        tbl.style = 'Table Grid'

        # Header row
        hdr_row = tbl.rows[0]
        for i, (h_text, w) in enumerate(zip(headers, col_widths_cm)):
            cell = hdr_row.cells[i]
            cell.width = Cm(w)
            cell_paragraph(cell, h_text, bold=True, size=9,
                           color=WHITE, align=WD_ALIGN_PARAGRAPH.CENTER)
            set_cell_bg(cell, INK)
            border = {'val': 'single', 'sz': 2, 'color': 'FFFFFF'}
            set_cell_border(cell,
                top=border, bottom=border, left=border, right=border)

        # Data rows
        for row_idx, row in enumerate(tc['rows']):
            tr      = tbl.add_row()
            bg      = LIGHT_GREY if row_idx % 2 == 0 else WHITE
            g_color = grade_color(row['grade'])

            data = [
                row['subject_name'],
                row['ca'],
                row['exam_score'],
                row['total'],
                row['grade'],
                row['remark'],
            ]
            aligns = [
                WD_ALIGN_PARAGRAPH.LEFT,
                WD_ALIGN_PARAGRAPH.CENTER,
                WD_ALIGN_PARAGRAPH.CENTER,
                WD_ALIGN_PARAGRAPH.CENTER,
                WD_ALIGN_PARAGRAPH.CENTER,
                WD_ALIGN_PARAGRAPH.CENTER,
            ]
            for i, (val, align) in enumerate(zip(data, aligns)):
                cell = tr.cells[i]
                cell.width = Cm(col_widths_cm[i])
                is_grade  = (i == 4)
                is_remark = (i == 5)
                cell_paragraph(
                    cell, val,
                    bold=(i == 0 or is_grade),
                    size=9,
                    color=g_color if (is_grade or is_remark) else INK,
                    align=align,
                )
                set_cell_bg(cell, bg)
                border = {'val': 'single', 'sz': 2, 'color': hex_color(MID_GREY)}
                set_cell_border(cell,
                    top=border, bottom=border, left=border, right=border)

        doc.add_paragraph()  # spacer

        # ── Remarks ───────────────────────────────────────────────────────────
        has_teacher   = bool(tc.get('teacher_comment', '').strip())
        has_principal = bool(tc.get('principal_comment', '').strip())

        if has_teacher or has_principal:
            add_paragraph(doc, 'Remarks', bold=True, size=11,
                          color=INK, space_before=6, space_after=4)

            rem_cols = 2 if (has_teacher and has_principal) else 1
            rem_tbl  = doc.add_table(rows=2, cols=rem_cols)
            rem_tbl.style = 'Table Grid'
            rem_w = Cm(8.5) if rem_cols == 2 else Cm(17.4)

            entries = []
            if has_teacher:
                entries.append(('Form Tutor', tc['teacher_comment']))
            if has_principal:
                entries.append(('Vice Principal', tc['principal_comment']))

            for col_i, (title, comment) in enumerate(entries):
                hc = rem_tbl.rows[0].cells[col_i]
                cc = rem_tbl.rows[1].cells[col_i]
                hc.width = rem_w
                cc.width = rem_w
                cell_paragraph(hc, title, bold=True, size=9, color=WHITE)
                cell_paragraph(cc, comment, bold=False, size=10,
                               color=DARK_GREY, italic=True)
                set_cell_bg(hc, INK)
                set_cell_bg(cc, LIGHT_GREY)
                border = {'val': 'single', 'sz': 2, 'color': 'FFFFFF'}
                for cell in (hc, cc):
                    set_cell_border(cell,
                        top=border, bottom=border, left=border, right=border)

            doc.add_paragraph()

        # ── Attendance summary ────────────────────────────────────────────────
        if tc['attendance'] is not None:
            add_paragraph(doc, 'Attendance Summary', bold=True, size=11,
                          color=INK, space_before=4, space_after=4)

            att_tbl = doc.add_table(rows=2, cols=3)
            att_tbl.style = 'Table Grid'
            att_labels = ['Days Present', 'Days Absent', 'Attendance Rate']
            att_values = [
                str(tc['days_present']),
                str(tc['days_absent']),
                f"{tc['attendance']}%",
            ]
            att_colors = [GREEN, RED, BLUE]
            att_w = Cm(5.8)

            for i, (lbl, val, col) in enumerate(
                    zip(att_labels, att_values, att_colors)):
                lc = att_tbl.rows[0].cells[i]
                vc = att_tbl.rows[1].cells[i]
                lc.width = att_w
                vc.width = att_w
                cell_paragraph(lc, lbl, bold=True, size=8, color=WHITE)
                cell_paragraph(vc, val, bold=True, size=14,
                               color=col, align=WD_ALIGN_PARAGRAPH.CENTER)
                set_cell_bg(lc, DARK_GREY)
                set_cell_bg(vc, LIGHT_GREY)
                border = {'val': 'single', 'sz': 2, 'color': 'FFFFFF'}
                for cell in (lc, vc):
                    set_cell_border(cell,
                        top=border, bottom=border, left=border, right=border)

            doc.add_paragraph()

        # ── Next term date ────────────────────────────────────────────────────
        if tc.get('next_term_begins'):
            p  = doc.add_paragraph()
            r1 = p.add_run('Next Term Begins: ')
            r1.bold = True
            r1.font.size = Pt(10)
            r1.font.color.rgb = DARK_GREY
            r2 = p.add_run(tc['next_term_begins'].strftime('%d %B %Y'))
            r2.font.size = Pt(10)
            r2.font.color.rgb = INK

        # ── Footer stamp ──────────────────────────────────────────────────────
        add_horizontal_rule(doc, color=GOLD)
        add_paragraph(doc, '✓  Result Verified & Approved — Evergreen Academy · Academic Board',
                      bold=False, size=9, color=DARK_GREY,
                      align=WD_ALIGN_PARAGRAPH.CENTER,
                      space_before=2, space_after=0)

    # ── Serialise to bytes ────────────────────────────────────────────────────
    buf = BytesIO()
    doc.save(buf)
    return buf.getvalue()


# ── Django view ───────────────────────────────────────────────────────────────

# def print_result_view(request, student, term_cards):
#     """
#     Call this from your main result view when the user clicks Print.
#     Example in urls.py:

#         path('result/print/', views.print_result, name='print_result'),

#     Example view wrapper in views.py:

#         def print_result(request):
#             student    = get_logged_in_student(request)
#             term_cards = build_term_cards(student)   # reuse your existing logic
#             return print_result_view(request, student, term_cards)
#     """
#     docx_bytes = build_result_docx(student, term_cards)
#     filename   = f"result_{student.student_id}_{student.full_name.replace(' ', '_')}.docx"

#     response = HttpResponse(
#         docx_bytes,
#         content_type='application/vnd.openxmlformats-officedocument.wordprocessingml.document',
#     )
#     response['Content-Disposition'] = f'attachment; filename="{filename}"'
#     return response


from docx2pdf import convert
import tempfile
import os

def print_result_view(request, student, term_cards):
    docx_bytes = build_result_docx(student, term_cards)
    # filename   = f"result_{student.student_id}_{student.full_name.replace(' ', '_')}"
    safe_id   = str(student.student_id).replace('/', '-').replace('\\', '-')
    safe_name = student.full_name.replace(' ', '_')
    filename  = f"result_{safe_id}_{safe_name}"

    with tempfile.TemporaryDirectory() as tmpdir:
        docx_path = os.path.join(tmpdir, f"{filename}.docx")
        pdf_path  = os.path.join(tmpdir, f"{filename}.pdf")

        with open(docx_path, 'wb') as f:
            f.write(docx_bytes)

        # Convert docx → pdf
        convert(docx_path, pdf_path)

        with open(pdf_path, 'rb') as f:
            pdf_bytes = f.read()

    response = HttpResponse(pdf_bytes, content_type='application/pdf')
    response['Content-Disposition'] = f'attachment; filename="{filename}.pdf"'
    return response


